import io

from app.domain.entities.document import Document
from docx import Document as DocxDocument

# Экспорт документа в Word

class DocxExporter:

    content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    file_extension = "docx"

    # Превращает документ в байты готового файла
    def export(self, document: Document) -> bytes:
        docx_doc = DocxDocument()
        docx_doc.add_heading(document.title, level=1)

        for line in document.content_snapshot.split("\n"):
            docx_doc.add_paragraph(line)

        buffer = io.BytesIO()
        docx_doc.save(buffer)
        return buffer.getvalue()